import colorama
from colorama import Fore
import time
import os
from datetime import datetime
from docx import Document
from doc_funcs import * 

# Changing the color of the cmd shell of the executable file
colorama.init()

# Welcome and doc choose
doc_model = model_choose(Fore.GREEN + 'Elije el numero de modelo a procesar')

# Set the path of the app to load the docs
doc_path = input_choose(Fore.GREEN + 'Dime el nombre de la carpeta. Pulsa S o di si para usar valor default "proc_doc\\docs"')
base_excel = input_choose(Fore.GREEN + 'Dime el nombre del excel de base a usar. Pulsa S o di si para usar valor default "Excel Modelos FAB.xslx"')
base_words = input_choose(Fore.GREEN + 'Dime el nombre del word de base a usar. Pulsa S o di si para usar valor default "cierre instruccion procedimiento.docx"')

if doc_path in ['Si', 'S', 'si', 's', 'sI']:
    doc_path = 'proc_doc\\docs'
else:
    doc_path = ''
if base_excel in ['Si', 'S', 'si', 's', 'sI']:
    base_excel = 'Excel Modelos FAB'
if base_words in ['Si', 'S', 'si', 's', 'sI']:
    base_words = 'cierre instruccion procedimiento'

# Path settings of doc and excel templates
current_dir = os.getcwd() + f"\\{doc_path}"
doc_path = current_dir + f"\\{base_words}.docx"
excel_path = current_dir + f"\\{base_excel}.xlsx"

# Reading Excel template
# TODO - PONER FALLO SI EL DOCUMENTO NO EXISTE
excel_dict = excel_to_json(excel_path)

# Functional transformations on the columns
tipo_a_anterior = str(int(float(excel_dict['NUMERO_A'][0])) - 1)
numero_a = str(int(float(excel_dict['NUMERO_A'][0])))
art_regla = str(int(float(excel_dict['ART_REGLA_ESCRUTINIO'][0])))
# Current date in Spanish legal format
now = datetime.now()
current_date = current_date_format(now)
# Correcting transormed values in the json dict
excel_dict['NUMERO_A'] = [numero_a]
excel_dict['A_ANTERIOR'] = [tipo_a_anterior]
excel_dict['CURRENT_DATE'] = [current_date]
excel_dict['ART_REGLA_ESCRUTINIO'] = [art_regla]

# Opening word doc
# TODO - PONER FALLO SI EL DOCUMENTO NO EXISTE
doc = Document(doc_path)

# Cleaning the dictionary
clean_dict = {}
for field, vals in excel_dict.items():
    try:
        list_vals = list(vals.values())
    except AttributeError:
        list_vals = vals
    try:
        list_vals.remove('nan')
    except ValueError:
        pass
    clean_dict[field] = list_vals

# Replacing the words in the doc
header = doc.sections[0].header
for field, vals in clean_dict.items():
    doc = iterate_paragraphs_and_headers(doc, doc.paragraphs, field, vals)
    doc = iterate_paragraphs_and_headers(doc, header.paragraphs, field, vals)

# Values for doc title
demandante = excel_dict['NOMBRE_DEMANDANTE'][0]
demandada = excel_dict['NOMBRE_DEMANDADA'][0]
tipo_a = excel_dict['TIPO_A'][0]

# Save the document
save_path = current_dir+f"\\{doc_model+' - '+tipo_a+' - '+demandante+' - '+demandada+' - '+current_date}.docx"

try:
    doc.save(save_path)
    print(f'Guardando el documento. Lo veras en la carpeta {current_dir}')
except Exception as e:
    print('No se puede guardar el documento por el error: ', e)
    time.sleep(5)

# Sleeping app to make user see the interface and final message
time.sleep(5)
