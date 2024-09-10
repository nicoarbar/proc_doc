import pandas as pd
from docx import Document
import streamlit as st

def current_date_format(date):
    months = ("Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre")
    day = date.day
    month = months[date.month - 1]
    year = date.year
    messsage = "{} de {} de {}".format(day, month, year)
    return messsage

def read_docx(file):
    doc = Document(file)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    return '\n'.join(content)

def streamlit_upload_csv(label, success, header_cols_list):
    doc_content = None
    doc = st.file_uploader(label)
    if doc is not None:
        doc_content = pd.read_csv(doc, names=header_cols_list)
        st.success(success)
    return doc_content

def streamlit_upload_docx(label, success):
    doc_content = None
    doc = st.file_uploader(label)
    if doc is not None:
        doc_content = read_docx(doc)
        st.success(success)
    return doc_content

def proc_doc_replace(doc, param):
    try:
        header = doc.sections[0].header
        n_params = len(param["Parameters"])
        for i in range(0, n_params):
            j = str(i)
            doc = iterate_paragraphs_and_headers(doc, doc.paragraphs, param["Parameters"][j], param["Value"][j])
            doc = iterate_paragraphs_and_headers(doc, header.paragraphs, param["Parameters"][j], param["Value"][j])
    except Exception as e:
        raise Exception(f'Doc is bad formatted by: {e}')
    return doc

def iterate_paragraphs_and_headers(doc, paragraphs, field, vals):
    for par in paragraphs:
        if field in par.text:
            if len(vals) > 1:
                par.text = par.text.replace(field, '\n'.join(vals))
            else:
                par.text = par.text.replace(field, vals[0])
    return doc