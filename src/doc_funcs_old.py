from pandas import *
import pandas as pd
import os
import time
import csv
from docx import Document
import streamlit as st

def model_choose(message):
    opt_dict = {
        '1': 'A1',
        '2': 'A2',
        '3': 'A3',
        '4': 'A4',
        '5': 'A5',
        '6': 'A6',
        '10': 'Todos los modelos'
    }
    for k, v in opt_dict.items():
        print(v, ': ', k)

    modelo = input(message)
    try:
        modelo_choose = opt_dict[modelo]
        print(f'Eljiste modelo {modelo_choose}')
    except KeyError:
        print(f'No existe el modelo {modelo}')
        modelo_choose = model_choose(message)
    return modelo_choose

def input_choose(message):
    input_choose = input(message)
    return input_choose

def excel_to_json(excel_path):
    xls = ExcelFile(excel_path)
    df = xls.parse(xls.sheet_names[0]).astype(str)
    excel_dict = df.to_dict()
    return excel_dict

def current_date_format(date):
    months = ("Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre")
    day = date.day
    month = months[date.month - 1]
    year = date.year
    messsage = "{} de {} de {}".format(day, month, year)
    return messsage

def doc_choose(current_dir, doc_name, doc_type):
    base_words = input_choose(Fore.GREEN + f'Dime el nombre del {doc_type} de base a usar. Pulsa Si para "{doc_name}"')
    if base_words in ['Si', 'S', 'si', 's', 'sI']:
        base_words = doc_name
    return current_dir + f"\\{base_words}"

def open_doc(current_dir, doc_name, doc_type, doc_func):
    # Path settings and reading Word template
    try:
        # Setting the path of the chosen doc
        doc_path = doc_choose(current_dir, doc_name, doc_type)
        # Checking if the doc exists
        if not os.path.isfile(doc_path):
            raise FileNotFoundError
        doc_object = doc_func(doc_path)
    except FileNotFoundError:
        print(f'No se encuentra el archivo Word que dices en la ruta: {doc_path}')
        doc_object = open_doc(current_dir, doc_name, doc_type, doc_func)
    except PermissionError:
        print(f'Primero cierra el documento que lo tienes abierto: {doc_path}')
        doc_object = open_doc(current_dir, doc_name, doc_type, doc_func)
    return doc_object

def raise_exception(message):
    print(message)
    time.sleep(5)
    raise Exception(message)

def num_to_str_fields(numeric_field):
    try:
        str_field = int(float(numeric_field))
    except Exception as e:
        raise_exception(f'Falla la limpieza del excel por el error: {e}')
    return str_field

def read_csv_as_rows(filename):
    with open(filename, mode='r') as file:
        csv_reader = csv.reader(file)
        for row in csv_reader:
            print(row)

def read_docx(file):
    #return Document(file)
    doc = Document(file)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    return '\n'.join(content)

def streamlit_upload_csv(label, success, header_cols_list):
    doc_content = None
    doc = st.file_uploader(label)
    if doc is not None:
        doc_content = pd.read_csv(doc, names=header_cols_list)
        st.success(success)
    return doc_content

def streamlit_upload_docx(label, success):
    doc_content = None
    doc = st.file_uploader(label)
    if doc is not None:
        doc_content = read_docx(doc)
        st.success(success)
    return doc_content

def proc_doc_replace(doc, param):
    try:
        n_params = len(param["Parameters"])
        for i in range(0, n_params):
            j = str(i)
            doc = iterate_paragraphs(doc, param["Parameters"][j], param["Value"][j])
            doc = iterate_headers(doc, param["Parameters"][j], param["Value"][j])
    except Exception as e:
        raise_exception(f'Doc is bad formatted by: {e}')
    return doc

def iterate_paragraphs(doc, field, vals):
    for par in doc.paragraphs:
        if field in par.text:
            if len(vals) > 1:
                par.text = par.text.replace(field, '\n'.join(vals))
            else:
                par.text = par.text.replace(field, vals[0])
    return doc

def iterate_headers(doc, field, vals):
    header = doc.sections[0].header
    for par in header.paragraphs:
        if field in par.text:
            if len(vals) > 1:
                par.text = par.text.replace(field, '\n'.join(vals))
            else:
                par.text = par.text.replace(field, vals[0])
    return doc